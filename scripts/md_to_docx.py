import markdown
from docx import Document
from docx.shared import Pt
import sys

def convert_md_to_docx(md_file, docx_file):
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)
    
    # Create a new Word document
    doc = Document()
    
    # Add the title
    title = "Документация системы обработки отчетов ЦРПТ"
    doc.add_heading(title, 0)
    
    # Add the converted content as paragraphs
    # (This is a very simplified approach - formatting will be basic)
    for line in html_content.split('<p>'):
        if '</p>' in line:
            text = line.split('</p>')[0]
            if text.strip():
                doc.add_paragraph(text)
    
    # Save the document
    doc.save(docx_file)
    print(f"Document saved as {docx_file}")

if __name__ == "__main__":
    md_file = "c:\\Users\\workb\\Downloads\\chzapi\\chzapiv3\\README.md"
    docx_file = "c:\\Users\\workb\\Downloads\\chzapi\\chzapiv3\\Documentation.docx"
    
    # Allow command line arguments
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
        
    convert_md_to_docx(md_file, docx_file)
